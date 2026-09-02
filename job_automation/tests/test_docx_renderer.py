"""Tests for output/docx_renderer.py"""
import pytest
from pathlib import Path
from docx import Document

from output.docx_renderer import (
    DocxRenderer,
    _classify_lines,
    _LineType,
    _sanitise_content,
    _validate_job_id,
    _validate_output_path,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
@pytest.fixture
def renderer(tmp_path):
    return DocxRenderer(output_dir=tmp_path / "docs")


# ---------------------------------------------------------------------------
# DocxRenderer — file creation
# ---------------------------------------------------------------------------
class TestDocxRendererFileCreation:
    def test_creates_output_directory(self, tmp_path):
        out_dir = tmp_path / "nested" / "docs"
        DocxRenderer(output_dir=out_dir)
        assert out_dir.exists()

    def test_generates_file_with_correct_name(self, renderer):
        path = renderer.render("job-abc-123", "My tailored resume content.")
        assert path.name == "job-abc-123.docx"
        assert path.exists()

    def test_generated_file_is_readable(self, renderer):
        content = "Jane Doe\nData Analyst\nExperience:\n- Wrote SQL queries"
        path = renderer.render("test-job", content)
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        assert "Jane Doe" in paragraphs
        assert "Data Analyst" in paragraphs

    def test_multiline_content_preserved(self, renderer):
        lines = ["Line 1", "", "Line 3", "Line 4"]
        path = renderer.render("multi-job", "\n".join(lines))
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert "Line 1" in texts
        assert "Line 3" in texts
        assert "Line 4" in texts

    def test_template_used_if_provided(self, tmp_path):
        template_path = tmp_path / "template.docx"
        template_doc = Document()
        template_doc.add_paragraph("TEMPLATE HEADER")
        template_doc.save(str(template_path))

        renderer = DocxRenderer(output_dir=tmp_path / "docs", template_path=template_path)
        path = renderer.render("tmpl-job", "Tailored content.")
        assert path.exists()

    def test_missing_template_falls_back_to_blank(self, tmp_path):
        renderer = DocxRenderer(
            output_dir=tmp_path / "docs",
            template_path=tmp_path / "nonexistent.docx",
        )
        path = renderer.render("fallback-job", "Some content.")
        assert path.exists()

    def test_invalid_template_extension_falls_back(self, tmp_path):
        renderer = DocxRenderer(
            output_dir=tmp_path / "docs",
            template_path=tmp_path / "bad.pdf",
        )
        path = renderer.render("ext-job", "Some content.")
        assert path.exists()


# ---------------------------------------------------------------------------
# DocxRenderer — resume formatting
# ---------------------------------------------------------------------------
class TestDocxRendererFormatting:
    def test_name_is_first_paragraph(self, renderer):
        path = renderer.render("fmt-name", "Alice Smith\nalice@email.com\n\nSUMMARY\nHello")
        doc = Document(str(path))
        non_empty = [p.text for p in doc.paragraphs if p.text.strip()]
        assert non_empty[0] == "Alice Smith"

    def test_section_headers_uppercased(self, renderer):
        path = renderer.render("fmt-section", "Name\n\nWork Experience\n- Did things")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert "WORK EXPERIENCE" in texts

    def test_bullet_text_without_leading_dash(self, renderer):
        path = renderer.render("fmt-bullet", "Name\n\nEXPERIENCE\n- Built pipelines")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert "Built pipelines" in texts
        assert "- Built pipelines" not in texts

    def test_role_line_detected(self, renderer):
        path = renderer.render(
            "fmt-role",
            "Name\n\nEXPERIENCE\nAcme Corp | Data Engineer | Jan 2022 – Present\n- Built things",
        )
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert "Acme Corp | Data Engineer | Jan 2022 – Present" in texts


# ---------------------------------------------------------------------------
# job_id validation
# ---------------------------------------------------------------------------
class TestValidateJobId:
    def test_valid_ids_pass(self):
        for jid in ("abc", "job-123", "job_456", "A1b2C3"):
            assert _validate_job_id(jid) == jid

    def test_path_traversal_rejected(self):
        for bad in ("../evil", "../../etc/passwd", "foo/bar", "foo\\bar"):
            with pytest.raises(ValueError):
                _validate_job_id(bad)

    def test_empty_string_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id("")

    def test_non_string_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id(None)  # type: ignore[arg-type]

    def test_leading_special_char_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id("-starts-with-dash")

    def test_too_long_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id("a" * 129)

    def test_null_byte_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id("job\x00id")

    def test_spaces_rejected(self):
        with pytest.raises(ValueError):
            _validate_job_id("job id")


# ---------------------------------------------------------------------------
# Content sanitisation
# ---------------------------------------------------------------------------
class TestSanitiseContent:
    def test_control_chars_stripped(self):
        result = _sanitise_content("hello\x00world\x07")
        assert "\x00" not in result
        assert "\x07" not in result
        assert "helloworld" in result

    def test_newlines_preserved(self):
        result = _sanitise_content("line1\nline2\nline3")
        assert result.count("\n") == 2

    def test_xml_tags_stripped(self):
        result = _sanitise_content("Hello <script>alert(1)</script> World")
        assert "<script>" not in result
        assert "Hello" in result
        assert "World" in result

    def test_markdown_stripped(self):
        result = _sanitise_content("**bold** and _italic_ and `code`")
        assert "**" not in result
        assert "_" not in result
        assert "`" not in result
        assert "bold" in result
        assert "italic" in result
        assert "code" in result

    def test_heading_markdown_stripped(self):
        result = _sanitise_content("## SKILLS\nPython, SQL")
        assert "##" not in result
        assert "SKILLS" in result

    def test_length_capped(self):
        long_text = "x" * 60_000
        result = _sanitise_content(long_text)
        assert len(result) <= 50_000

    def test_non_string_raises_type_error(self):
        with pytest.raises(TypeError):
            _sanitise_content(12345)  # type: ignore[arg-type]

    def test_windows_line_endings_normalised(self):
        result = _sanitise_content("line1\r\nline2\r\nline3")
        assert "\r" not in result
        assert result.count("\n") == 2


# ---------------------------------------------------------------------------
# Output path traversal guard
# ---------------------------------------------------------------------------
class TestValidateOutputPath:
    def test_normal_path_allowed(self, tmp_path):
        path = _validate_output_path(tmp_path, "job-123.docx")
        assert path == (tmp_path / "job-123.docx").resolve()

    def test_traversal_rejected(self, tmp_path):
        with pytest.raises(ValueError, match="escapes"):
            _validate_output_path(tmp_path, "../outside.docx")

    def test_sibling_directory_with_shared_prefix_rejected(self, tmp_path):
        """
        A plain str.startswith(output_dir) check would wrongly let this
        through: "<tmp_path>docs2/x.docx" starts with "<tmp_path>docs" as a
        string even though docs2 is a sibling directory, not a child of docs.
        """
        docs_dir = tmp_path / "docs"
        docs_dir.mkdir()
        with pytest.raises(ValueError, match="escapes"):
            _validate_output_path(docs_dir, "../docs2/x.docx")


# ---------------------------------------------------------------------------
# Line classifier
# ---------------------------------------------------------------------------
class TestClassifyLines:
    def test_first_line_is_name(self):
        lines = _classify_lines(["John Doe", "john@email.com"])
        assert lines[0].kind == _LineType.NAME

    def test_email_line_is_contact(self):
        lines = _classify_lines(["John Doe", "john@email.com"])
        assert lines[1].kind == _LineType.CONTACT

    def test_all_caps_section_detected(self):
        lines = _classify_lines(["Name", "", "WORK EXPERIENCE"])
        kinds = [l.kind for l in lines]
        assert _LineType.SECTION in kinds

    def test_known_section_keyword_detected(self):
        lines = _classify_lines(["Name", "", "Education"])
        kinds = [l.kind for l in lines]
        assert _LineType.SECTION in kinds

    def test_bullet_detected(self):
        lines = _classify_lines(["Name", "", "SKILLS", "- Python"])
        assert lines[-1].kind == _LineType.BULLET

    def test_role_line_detected(self):
        lines = _classify_lines(["Name", "", "EXPERIENCE", "Acme | Engineer | Jan 2022"])
        assert lines[-1].kind == _LineType.ROLE

    def test_bullet_containing_month_substring_not_misclassified_as_role(self):
        """
        Regression test: bullets starting with "-" that happen to contain a
        3-letter month-abbreviation substring (Django->jan, innovative->nov,
        august/augment->aug, december/doctor->dec...) must stay BULLET, not
        get misclassified as ROLE (which used to skip dash-stripping and
        left "-  " literally in the rendered text).
        """
        bullets = [
            "- Shipped a Django-based platform used by four customers.",
            "- Built an innovative pipeline that cut review time significantly.",
            "- Automated deployments using augmented monitoring tooling.",
            "- Delivered as the on-call doctor for the platform's SLA program.",
        ]
        for b in bullets:
            lines = _classify_lines(["Name", "", "SECTION", b])
            assert lines[-1].kind == _LineType.BULLET, f"misclassified: {b!r}"
            assert not lines[-1].text.startswith("-")

    def test_role_line_without_pipe_not_matched(self):
        """A bullet with a bare hyphen and a date-like year should not be
        mistaken for a role line just because it has a year in it."""
        lines = _classify_lines(["Name", "", "SECTION", "- Migrated 2 legacy 2019-era services."])
        assert lines[-1].kind == _LineType.BULLET

    def test_empty_line_classified(self):
        lines = _classify_lines(["Name", ""])
        assert lines[1].kind == _LineType.EMPTY

    def test_body_fallthrough(self):
        lines = _classify_lines(["Name", "", "SUMMARY", "A professional with experience."])
        assert lines[-1].kind == _LineType.BODY
