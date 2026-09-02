"""
output/docx_renderer.py — Renders tailored resume text to .docx files.

Safe write workflow
-------------------
1. Sanitise input text and validate job_id.
2. Write to a temp file (.tmp.docx) in the output directory.
3. Re-open and validate the temp file (non-empty, required sections).
4. Atomically rename temp → final path only after validation passes.
5. On failure: retain temp for debugging (within retention window), raise
   DocxValidationError so the caller can set DOCX_GENERATION_FAILED.

Smart resume formatting
-----------------------
Lines are classified before writing so the DOCX has structure:
  - NAME    → bold, 14pt
  - CONTACT → 10pt, grey-ish
  - SECTION → bold, 12pt, all-caps
  - ROLE    → italic, 11pt (employer | title | dates pattern)
  - BULLET  → 11pt, leading dash stripped
  - BODY    → 11pt
  - EMPTY   → spacer paragraph

Security
--------
  - _validate_job_id() rejects path-traversal characters, null bytes, spaces
  - _validate_output_path() ensures the resolved final path stays inside output_dir
  - _sanitise_content() strips control chars, XML tags, markdown syntax, caps length
"""
from __future__ import annotations

import enum
import logging
import re
import shutil
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

DEFAULT_OUTPUT_DIR = Path("output/docs")
DEFAULT_TEMP_RETENTION_HOURS = 24
_MAX_JOB_ID_LEN = 128
_MAX_CONTENT_LEN = 50_000

# Compact page layout — a blank python-docx Document() inherits Word's default
# template, which sets 1.25"/1.0" margins and (buried in styles.xml
# docDefaults, not visible via python-docx's paragraph_format API) 10pt of
# space AFTER EVERY PARAGRAPH plus 1.15 line spacing. Across a ~60-paragraph
# resume that default alone is ~8 inches of pure inter-paragraph whitespace —
# the dominant reason an otherwise normal-length resume renders at 4 pages.
# These constants override both the margins and the per-paragraph spacing
# explicitly so the page count reflects the actual content, not template
# defaults nobody asked for.
_PAGE_MARGIN_INCHES = 0.5

# (space_before_pt, space_after_pt, line_spacing, font_pt) per _LineType.
_LINE_FORMAT: dict[str, tuple[float, float, float, float]] = {
    "name": (0, 1, 1.0, 15),
    "contact": (0, 5, 1.0, 9.5),
    "section": (5, 2, 1.0, 11),
    "role": (2, 0, 1.0, 10),
    "bullet": (0, 1, 1.0, 10),
    "body": (0, 2, 1.0, 10),
    "empty": (0, 0, 1.0, 3),  # near-zero-height spacer, not another 10pt gap
}

# Known section header keywords (title-cased forms; matched case-insensitively)
_SECTION_KEYWORDS = frozenset({
    "experience", "work experience", "education", "skills", "summary",
    "professional summary", "objective", "certifications", "projects",
    "achievements", "awards", "languages", "interests", "references",
    "employment", "qualifications", "profile",
})

# Regex that identifies a role/employer line: "Company | Title | Mon YYYY - Mon YYYY".
# Requires the pipe delimiter the output-format contract mandates (see
# ai/prompts.py _OUTPUT_FORMAT), plus a year or "present".
#
# NOTE: a previous version matched on any "|" OR "-" (not just "|"), followed
# anywhere later in the line by a month-abbreviation substring. Since every
# bullet line starts with "-", and 3-letter month abbreviations are common
# word fragments (Django contains "jan", innovative/november contain "nov",
# august/augment contain "aug", december/doctor contain "dec"...), that
# matched a large fraction of ordinary bullets as ROLE lines — which skips
# the bullet dash-stripping, so they rendered as "-  Some bullet text" in
# italics instead of a proper bullet. Requiring "|" plus an actual year (the
# format the AI is instructed to always emit) removes the false positives.
_ROLE_RE = re.compile(r"\|.*(?:\d{4}|\bpresent\b)", re.IGNORECASE)

# Job-id allowed characters: alphanumeric, dash, underscore; must start with alnum
_JOB_ID_RE = re.compile(r'^[A-Za-z0-9][A-Za-z0-9_\-]{0,127}$')

# Control characters to strip, excluding LF (\x0a) and TAB (\x09).
# NOTE: a previous version of this pattern was `[^\S\n\t]|[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]`.
# `[^\S\n\t]` is the negation of "non-whitespace, LF, TAB" — i.e. it matches
# ALL OTHER WHITESPACE, including the literal space character, so it stripped
# every space out of every resume. Fixed to only match actual control chars.
_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

# Strip XML/HTML tags
_XML_TAG_RE = re.compile(r"<[^>]+>")
# Strip markdown: bold (**), italic (single _), inline code (`)
_MARKDOWN_RE = re.compile(r"(\*\*|__|\*|_|`|#+\s*)")


# ---------------------------------------------------------------------------
# Public types
# ---------------------------------------------------------------------------

class DocxValidationError(Exception):
    """Raised when a generated DOCX fails the content validation gate."""


class _LineType(enum.Enum):
    NAME = "name"
    CONTACT = "contact"
    SECTION = "section"
    ROLE = "role"
    BULLET = "bullet"
    BODY = "body"
    EMPTY = "empty"


class _ClassifiedLine:
    __slots__ = ("text", "kind")

    def __init__(self, text: str, kind: _LineType) -> None:
        self.text = text
        self.kind = kind


# ---------------------------------------------------------------------------
# Security helpers (public for test import)
# ---------------------------------------------------------------------------

def _validate_job_id(job_id: object) -> str:
    """
    Validate and return the job_id string.

    Raises ValueError for:
    - Non-string input
    - Empty string
    - Path-traversal characters (/, \\, ..)
    - Null bytes or spaces
    - Leading special characters (dash, underscore)
    - Length > 128 characters
    """
    if not isinstance(job_id, str):
        raise ValueError(f"job_id must be a str, got {type(job_id).__name__!r}")
    if not job_id:
        raise ValueError("job_id must not be empty.")
    if "\x00" in job_id:
        raise ValueError("job_id contains a null byte.")
    if " " in job_id:
        raise ValueError("job_id must not contain spaces.")
    if "/" in job_id or "\\" in job_id:
        raise ValueError("job_id must not contain path separators.")
    if ".." in job_id:
        raise ValueError("job_id must not contain '..'.")
    if not _JOB_ID_RE.match(job_id):
        raise ValueError(
            f"job_id {job_id!r} is invalid. Must start with alphanumeric and "
            "contain only letters, digits, dashes, or underscores (max 128 chars)."
        )
    return job_id


def _validate_output_path(output_dir: Path, filename: str) -> Path:
    """
    Resolve the full output path and verify it stays inside output_dir.

    Raises ValueError if the resolved path escapes the output directory.
    """
    resolved_dir = Path(output_dir).resolve()
    resolved_path = (resolved_dir / filename).resolve()
    # NOTE: a plain str.startswith(resolved_dir) check is a classic bypass —
    # "/out/docs2/x" starts with "/out/docs" even though it's a sibling
    # directory, not a child. Use real path containment instead.
    if resolved_path != resolved_dir and resolved_dir not in resolved_path.parents:
        raise ValueError(
            f"Output path {resolved_path!r} escapes output directory {resolved_dir!r}."
        )
    return resolved_path


def _sanitise_content(text: object) -> str:
    """
    Sanitise resume text before writing to DOCX.

    Raises TypeError for non-string input.

    Steps:
    1. Normalise Windows line endings to LF.
    2. Strip control characters (except LF/TAB).
    3. Strip XML/HTML tags.
    4. Strip markdown syntax (**, *, _, `, #).
    5. Cap length at _MAX_CONTENT_LEN characters.
    """
    if not isinstance(text, str):
        raise TypeError(f"content must be a str, got {type(text).__name__!r}")

    # Normalise Windows line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Strip control characters except LF (\n) and TAB (\t)
    text = _CONTROL_CHAR_RE.sub("", text)

    # Strip XML/HTML tags
    text = _XML_TAG_RE.sub("", text)

    # Strip markdown syntax
    text = _MARKDOWN_RE.sub("", text)

    # Cap length
    if len(text) > _MAX_CONTENT_LEN:
        text = text[:_MAX_CONTENT_LEN]

    return text


# ---------------------------------------------------------------------------
# Line classifier (public for test import)
# ---------------------------------------------------------------------------

def _classify_lines(lines: list[str]) -> list[_ClassifiedLine]:
    """
    Classify a list of text lines into _LineType categories.

    Classification rules (evaluated in order):
    1. Empty line → EMPTY
    2. First non-empty line → NAME
    3. Line contains '@' with no spaces → CONTACT
    4. All-caps line (≥2 words or a known keyword) → SECTION
    5. Known section keyword (title-cased) → SECTION
    6. Matches role pattern (contains '|' or '–' with date-ish text) → ROLE
    7. Starts with '-' or '•' → BULLET
    8. Otherwise → BODY
    """
    result: list[_ClassifiedLine] = []
    name_seen = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            result.append(_ClassifiedLine("", _LineType.EMPTY))
            continue

        # NAME: first non-empty line
        if not name_seen:
            name_seen = True
            result.append(_ClassifiedLine(stripped, _LineType.NAME))
            continue

        # CONTACT: contains '@' (email) or looks like a phone/URL contact
        if "@" in stripped and " " not in stripped:
            result.append(_ClassifiedLine(stripped, _LineType.CONTACT))
            continue

        # SECTION: all-caps (letters only), or known keyword
        words_only = re.sub(r"[^A-Za-z ]", "", stripped)
        if words_only.strip() and words_only.strip() == words_only.strip().upper():
            result.append(_ClassifiedLine(stripped.upper(), _LineType.SECTION))
            continue

        if stripped.lower() in _SECTION_KEYWORDS:
            result.append(_ClassifiedLine(stripped.upper(), _LineType.SECTION))
            continue

        # ROLE: employer | title | dates
        if _ROLE_RE.search(stripped):
            result.append(_ClassifiedLine(stripped, _LineType.ROLE))
            continue

        # BULLET
        if stripped.startswith(("-", "•", "*")):
            bullet_text = re.sub(r"^[-•*]\s*", "", stripped)
            result.append(_ClassifiedLine(bullet_text, _LineType.BULLET))
            continue

        # Default: BODY
        result.append(_ClassifiedLine(stripped, _LineType.BODY))

    return result


# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------

class DocxRenderer:
    """
    Writes plain-text tailored resumes to DOCX files with:
    - Smart line classification and formatting
    - Input sanitisation and job_id validation
    - Atomic write-validate-move pipeline
    - Idempotent rendering (skips valid existing files)

    Args:
        output_dir:             directory where final files are saved.
        template_path:          optional path to a .docx template file.
        required_sections:      list of strings that must appear (case-insensitive)
                                in the DOCX body text.  Empty → only non-empty body required.
        temp_retention_hours:   how long failed-validation temp files are kept.
    """

    def __init__(
        self,
        output_dir: Path = DEFAULT_OUTPUT_DIR,
        template_path: Optional[Path] = None,
        required_sections: Optional[list[str]] = None,
        temp_retention_hours: int = DEFAULT_TEMP_RETENTION_HOURS,
    ) -> None:
        self._output_dir = Path(output_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)
        self._template_path = (
            Path(template_path)
            if template_path and Path(template_path).suffix.lower() == ".docx"
            else None
        )
        self._required_sections: list[str] = [
            s.lower() for s in (required_sections or [])
        ]
        self._temp_retention_hours = temp_retention_hours

    # ------------------------------------------------------------------
    # Public
    # ------------------------------------------------------------------

    def render(self, job_id: str, tailored_text: str) -> Path:
        """
        Sanitise, write, validate, and atomically move a DOCX.

        Returns the final Path on success.
        Raises DocxValidationError on content validation failure.
        Raises IOError on unexpected write errors.
        """
        job_id = _validate_job_id(job_id)
        safe_text = _sanitise_content(tailored_text)

        final_path = _validate_output_path(self._output_dir, f"{job_id}.docx")

        # Idempotency: return existing valid file
        if final_path.exists():
            try:
                self._validate(final_path)
                logger.info("DOCX already valid — skipping re-render: %s", final_path)
                return final_path
            except DocxValidationError:
                logger.warning("Existing DOCX failed validation — re-rendering: %s", final_path)

        tmp_path = _validate_output_path(self._output_dir, f"{job_id}.tmp.docx")

        try:
            doc = self._load_template_or_new()
            self._write_content(doc, safe_text)
            doc.save(str(tmp_path))
            logger.debug("DOCX temp written: %s", tmp_path)

            self._validate(tmp_path)

            shutil.move(str(tmp_path), str(final_path))
            logger.info("DOCX ready: %s", final_path)
            return final_path

        except DocxValidationError:
            logger.warning(
                "DOCX validation failed for job %s — temp retained at %s (for %dh)",
                job_id, tmp_path, self._temp_retention_hours,
            )
            raise

        except Exception as exc:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)
            raise IOError(f"DOCX write failed for job {job_id}: {exc}") from exc

    def cleanup_stale_temps(self) -> int:
        """
        Delete .tmp.docx files older than temp_retention_hours.
        Returns the count of files deleted.
        """
        cutoff = datetime.now(timezone.utc) - timedelta(hours=self._temp_retention_hours)
        deleted = 0
        for tmp_file in self._output_dir.glob("*.tmp.docx"):
            mtime = datetime.fromtimestamp(tmp_file.stat().st_mtime, tz=timezone.utc)
            if mtime < cutoff:
                tmp_file.unlink(missing_ok=True)
                logger.debug("Deleted stale temp DOCX: %s", tmp_file)
                deleted += 1
        if deleted:
            logger.info("Cleaned up %d stale temp DOCX file(s).", deleted)
        return deleted

    # ------------------------------------------------------------------
    # Internals
    # ------------------------------------------------------------------

    def _load_template_or_new(self) -> Document:
        if self._template_path and self._template_path.exists():
            try:
                return Document(str(self._template_path))
            except Exception:
                logger.warning("Template load failed — falling back to blank document.")
        return Document()

    def _write_content(self, doc: Document, text: str) -> None:
        """Classify lines and write with appropriate, page-count-conscious formatting."""
        if self._template_path is None:
            for section in doc.sections:
                section.left_margin = Inches(_PAGE_MARGIN_INCHES)
                section.right_margin = Inches(_PAGE_MARGIN_INCHES)
                section.top_margin = Inches(_PAGE_MARGIN_INCHES)
                section.bottom_margin = Inches(_PAGE_MARGIN_INCHES)

        lines = text.splitlines()
        classified = _classify_lines(lines)

        for cl in classified:
            para = doc.add_paragraph()
            space_before, space_after, line_spacing, font_pt = _LINE_FORMAT[cl.kind.value]
            pf = para.paragraph_format
            pf.space_before = Pt(space_before)
            pf.space_after = Pt(space_after)
            pf.line_spacing = line_spacing

            if cl.kind == _LineType.EMPTY:
                continue  # blank spacer line — height already minimised via _LINE_FORMAT

            run = para.add_run(cl.text)
            run.font.size = Pt(font_pt)

            if cl.kind == _LineType.NAME:
                run.bold = True

            elif cl.kind == _LineType.CONTACT:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            elif cl.kind == _LineType.SECTION:
                run.bold = True

            elif cl.kind == _LineType.ROLE:
                run.italic = True

            elif cl.kind == _LineType.BULLET:
                para.style = doc.styles["List Bullet"] if "List Bullet" in [s.name for s in doc.styles] else para.style
                # Re-apply spacing: assigning a built-in style resets paragraph_format.
                pf = para.paragraph_format
                pf.space_before = Pt(space_before)
                pf.space_after = Pt(space_after)
                pf.line_spacing = line_spacing

            # else: BODY — font size already set above, no extra styling

    def _validate(self, path: Path) -> None:
        """
        Validate a DOCX file for readability and content.
        Raises DocxValidationError on failure.
        """
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise DocxValidationError(f"DOCX could not be opened: {exc}") from exc

        body_text = "\n".join(p.text for p in doc.paragraphs)
        if not body_text.strip():
            raise DocxValidationError("DOCX body is empty — no paragraphs found.")

        body_lower = body_text.lower()
        missing = [s for s in self._required_sections if s not in body_lower]
        if missing:
            raise DocxValidationError(
                f"DOCX is missing required section(s): {missing}"
            )


def build_renderer_from_config(config: dict) -> DocxRenderer:
    """Convenience factory from config dict."""
    out_cfg = config.get("output", {})
    output_dir = Path(out_cfg.get("docx_output_dir", "output/docs"))
    template_path_str = out_cfg.get("docx_template_path", "")
    template_path = Path(template_path_str) if template_path_str else None
    required_sections: list[str] = out_cfg.get("docx_validation_required_sections", [])
    temp_retention_hours: int = int(
        out_cfg.get("docx_temp_retention_hours", DEFAULT_TEMP_RETENTION_HOURS)
    )
    return DocxRenderer(
        output_dir=output_dir,
        template_path=template_path,
        required_sections=required_sections,
        temp_retention_hours=temp_retention_hours,
    )
